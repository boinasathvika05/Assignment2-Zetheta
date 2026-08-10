import os
import re

try:
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    os.system("pip install python-docx")
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def convert_md_file_to_docx(md_path, docx_path):
    print(f"Converting {os.path.basename(md_path)} -> {os.path.basename(docx_path)}...")
    doc = docx.Document()
    
    # Configure document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []

    def flush_table():
        nonlocal in_table, table_rows
        if not table_rows:
            return
        
        # Determine cols
        header_cols = [c.strip() for c in table_rows[0].split('|')[1:-1]]
        num_cols = len(header_cols)
        if num_cols == 0:
            table_rows = []
            in_table = False
            return

        tbl = doc.add_table(rows=0, cols=num_cols)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header row
        hdr_cells = tbl.add_row().cells
        for idx, text in enumerate(header_cols):
            hdr_cells[idx].text = text
            set_cell_background(hdr_cells[idx], "1A365D") # Navy blue header
            for p in hdr_cells[idx].paragraphs:
                for run in p.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.name = "Segoe UI"
                    run.font.size = Pt(10)

        # Data rows
        for r in table_rows[1:]:
            if "---" in r or "===" in r:
                continue
            cols = [c.strip() for c in r.split('|')[1:-1]]
            if len(cols) != num_cols:
                continue
            row_cells = tbl.add_row().cells
            for idx, text in enumerate(cols):
                row_cells[idx].text = text
                for p in row_cells[idx].paragraphs:
                    for run in p.runs:
                        run.font.name = "Segoe UI"
                        run.font.size = Pt(9.5)

        doc.add_paragraph() # Spacing
        table_rows = []
        in_table = False

    for line in lines:
        raw_line = line.rstrip("\n")

        # Code block handling
        if raw_line.strip().startswith("```"):
            if in_code_block:
                # Flush code block
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(40, 40, 40)
                code_lines = []
                in_code_block = False
            else:
                if in_table:
                    flush_table()
                in_code_block = True
                code_lines = []
            continue

        if in_code_block:
            code_lines.append(raw_line)
            continue

        # Table handling
        if "|" in raw_line and ("---" in raw_line or len(raw_line.split("|")) > 2):
            in_table = True
            table_rows.append(raw_line)
            continue
        elif in_table:
            flush_table()

        # Headings
        if raw_line.startswith("# "):
            h = doc.add_heading(level=1)
            run = h.add_run(raw_line[2:].strip())
            run.font.name = "Segoe UI"
            run.font.color.rgb = RGBColor(26, 54, 93) # Deep Navy
            run.font.bold = True
        elif raw_line.startswith("## "):
            h = doc.add_heading(level=2)
            run = h.add_run(raw_line[3:].strip())
            run.font.name = "Segoe UI"
            run.font.color.rgb = RGBColor(43, 108, 176) # Slate Blue
            run.font.bold = True
        elif raw_line.startswith("### "):
            h = doc.add_heading(level=3)
            run = h.add_run(raw_line[4:].strip())
            run.font.name = "Segoe UI"
            run.font.color.rgb = RGBColor(45, 55, 72)
            run.font.bold = True
        elif raw_line.startswith("#### "):
            h = doc.add_heading(level=4)
            run = h.add_run(raw_line[5:].strip())
            run.font.name = "Segoe UI"
            run.font.bold = True
        elif raw_line.strip().startswith("- ") or raw_line.strip().startswith("* "):
            p = doc.add_paragraph(style='List Bullet')
            text = raw_line.strip()[2:]
            run = p.add_run(text)
            run.font.name = "Segoe UI"
            run.font.size = Pt(10.5)
        elif raw_line.strip() == "":
            pass
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(raw_line)
            run.font.name = "Segoe UI"
            run.font.size = Pt(10.5)

    if in_table:
        flush_table()

    doc.save(docx_path)
    print(f"Saved: {docx_path}")

def main():
    workspace = r"c:\Users\Sathvika\Downloads\zetheta assign 2"
    md_files = [
        "README.md",
        "USER_GUIDE.md",
        "DEPLOYMENT_GUIDE.md",
        "ARCHITECTURE.md",
        "REQUIREMENT_TRACEABILITY_MATRIX.md",
        "CHANGELOG.md"
    ]
    for md_file in md_files:
        full_md_path = os.path.join(workspace, md_file)
        if os.path.exists(full_md_path):
            docx_file = md_file.replace(".md", ".docx")
            full_docx_path = os.path.join(workspace, docx_file)
            convert_md_file_to_docx(full_md_path, full_docx_path)

if __name__ == "__main__":
    main()
